from googleapiclient.discovery import build
from google.oauth2 import service_account
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
import os
import matplotlib.pyplot as plt
from tkinter import *
from PIL import ImageTk, Image

root = Tk()
root.title('Automated Turtle Report Generator')

title_lab = Label(root, bg='white',
                  text='Welcome to the automated turtle report generator.\nJust select your report start date and end date, '
                             ' then click generate\nreport. The report will be generated as a word doc in the same'
                             ' folder as this\nprogram. Please input the date in this format: YYYY-mm-dd. Enjoy!!!!!')
title_lab.pack()


logo = ImageTk.PhotoImage(Image.open('Turtle_logo.jpg'))
panel = Label(root, image=logo).pack()

root['bg'] = 'white'

start_lab = Label(root, text='Select Start Date', bg='white')
start_lab.pack()

start_entry = Entry(root)
start_entry.pack()

end_lab = Label(root, text='Select End Date', bg='white')
end_lab.pack()

end_entry = Entry(root)
end_entry.pack()

def generate():

    start_date = str('\'' + start_entry.get() + '\'')
    end_date = str('\'' + end_entry.get() + '\'')

    SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
    SERVICE_ACCOUNT_FILE = 'keys.json'

    creds = None
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)

    # The ID spreadsheet.
    SAMPLE_SPREADSHEET_ID = '1X3rJLFptPcdWX2uIrAco9WKHy2Nv4iKY08cnq-vJQ-U'

    service = build('sheets', 'v4', credentials=creds)

    # Call the Sheets API
    sheet = service.spreadsheets()
    result = sheet.values().get(spreadsheetId=SAMPLE_SPREADSHEET_ID,
                                range='Sheet1!A1:BI3000').execute()
    values = result.get('values', [])
#    print(values)

    df = pd.DataFrame(values, columns=values[0])
    df['year'].astype(str)


#    df = pd.read_excel('Turtle_database.xlsx')
    df = df.replace(np.nan, '', regex=True)
    df['all.tags'] = df['remigrant'] + df['reclutch'] + df['new.tag']

    for u in range(0, len(df)):
        if 'est' in df['action'][u]:
            df['action'][u] = 'Nest'
        else:
            df['action'][u] = df['action'][u]

    for v in range(0, len(df)):
        if 'Fuw' in df['Location'][v]:
            df['Location'][v] = 'Fuwairit'
        else:
            df['Location'][v] = df['Location'][v]

    for p in range(0, len(df)):
        if 'Laf' in df['Location'][p]:
            df['Location'][p] = 'Ras Laffan'
        else:
            df['Location'][p] = df['Location'][p]

    for k in range(0, len(df)):
        if 'Gha' in df['Location'][k]:
            df['Location'][k] = 'Al Ghariyah'
        else:
            df['Location'][k] = df['Location'][k]



    this_year = df[df['year'] == start_date[1:5]]
    this_year['date'] = pd.to_datetime(this_year['nest date'])
    this_year.sort_values(by='date')

    nth = {
        '1': "first",
        '2': "second",
        '3': "third",
        '4': "fourth",
        '5': "fifth"
    }

    mask = (this_year['date'] >= start_date) & (this_year['date'] <= end_date)
    period = this_year.loc[mask]
    per = period.replace(np.nan, '', regex=True)

    surveys = per['date'].unique()

    pics = os.listdir('weekly report pics')

    new_mask = (this_year['date'] >= pd.to_datetime(str(start_date[1:5] + '-04-01'))) & (this_year['date'] <= end_date)
    all_period = this_year.loc[new_mask]
    all_counts = all_period.groupby(['action', 'Location']).agg('count').reset_index()
    all_nests = all_counts[all_counts['action'] == 'Nest']

    period_counts = per.groupby(['action', 'Location']).agg('count').reset_index()
    nest_counts = period_counts[period_counts['action'] == 'Nest']
    fca_counts = period_counts[period_counts['action'] == 'FCA']
    fcu_counts = period_counts[period_counts['action'] == 'FCU']

    # Pie chart, where the slices will be ordered and plotted counter-clockwise:
    labels = all_nests['Location']
    sizes = all_nests['nest date']
    #explode = (0, 0.1, 0, 0)  # only "explode" the 2nd slice (i.e. 'Hogs')

    fig1, ax1 = plt.subplots()
    ax1.pie(sizes, labels=labels, autopct='%1.1f%%',
            shadow=True, startangle=90)
    ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
    plt.savefig('pie.jpg')


    cum_count = all_period.groupby(['action', 'date']).agg('count').reset_index()
    nest_cum_count = cum_count[cum_count['action'] == 'Nest']
    nest_cum_count

    Date = nest_cum_count['date']
    nest_freq = nest_cum_count['nest date']

    fig = plt.figure()
    plt.bar(Date, nest_freq)
    plt.title('')
    plt.xlabel('Nest date', fontsize=14)
    plt.ylabel('Nest frequency', fontsize=14)
    plt.xticks(rotation=90)
    plt.savefig('Nest_freq.jpg', dpi=fig.dpi)

    cum_tot = []
    nest_arr = np.array(nest_freq)
    for i in range(0, (len(nest_freq)-1)):
        cum_tot.append(sum(nest_arr[0:i]))
        i+=1

    cum_date = Date[1:len(Date)]

    fig = plt.figure()
    plt.plot(cum_date, cum_tot)
    plt.title('')
    plt.xlabel('Nest date', fontsize=14)
    plt.ylabel('Cumulative nests', fontsize=14)
    plt.xticks(rotation=90)
    plt.savefig('Cumulative_nests.jpg', dpi=fig.dpi)


    tab_dat = all_period.groupby(['action', 'Location']).agg('count').reset_index()
    tab_df = tab_dat[tab_dat['action'] == 'Nest']
    tab_locs = np.array(tab_df['Location'])
    tab_nests = np.array(tab_df['nest date'])
    records = []
    for k in range(0, len(tab_locs)):
        records.append(
            [tab_locs[k], tab_nests[k]]
        )



    ######################################################################################################################################################

            #       BUILDING DOCUMENT       #

    ######################################################################################################################################################

    document = Document()

    document.add_heading('Marine Turtle Conservation and Monitoring in Ras Laffan Industrial City and other sites in the State of Qatar', 0)

    document.add_heading('Bi-Weekly Report: ' + str(start_date)[1:-1] + ' to ' + str(end_date)[1:-1], 1)
    document.add_heading('Prepared by the Environmental Science Centre of Qatar University for the 2021 season', 1)

    document.add_paragraph('')

    document.add_picture('Logos.png', width=Inches(5.5))

    document.add_page_break()

    document.add_heading('Surveys Conducted', level=1)

    for j in range(0, len(surveys)):

        surveys.sort()
        per.sort_values(by='date')
        new = per[per['date'] == surveys[j]]
        document.add_heading(str(surveys[j])[0:10], level=2)

        for i in range(0, len(new)):

            if (new.iloc[i]['action'] == 'Nest') & (new.iloc[i]['Location'] == 'Ras Rakkan'):
                    document.add_paragraph(str(len(new)) + ' nests were found on Ras Rakkan on this survey.', style='List Bullet')
                    break

            elif (new.iloc[i]['action'] == 'Nest') & (new.iloc[i]['Location'] == 'Umm Tais'):
                    document.add_paragraph(str(len(new)) + ' nests were found on Umm Tais on this survey.', style='List Bullet')
                    break

            if (new.iloc[i]['action'] == 'Nest') & (new.iloc[i]['remigrant'] != ''):
                    document.add_paragraph('A live turtle nesting was observed on ' + new.iloc[i]['Location'] + '. This turtle was a remigrant that was previously fitted with tag nos. ' + new.iloc[i]['remigrant'] + ' in ' + str(df[df['all.tags'].str.contains(new.iloc[i]['remigrant'][2:6])].tail(2)['Location'].values[0]) + '. She was previously captured nesting ' + str(new.iloc[i]['remigrant.period']) + ' years ago on ' + str(df[df['all.tags'].str.contains(new.iloc[i]['remigrant'][2:6])].tail(2)['nest date'].values[0]) + '. The clutch she deposited contained ' + str(new.iloc[i]['eggs']) + ' eggs which were relocated to the ' + str(new.iloc[i]['hatchery site']) + ' hatchery as nest number ' + str(new.iloc[i]['hatchery nest number']) + '. She had a curved carapace length (CCL) of ' + str(float(new.iloc[i]['ccl'])) + ' and a curved carapace width (CCW) of ' + str(float(new.iloc[i]['ccw'])) + '.', style='List Bullet')

            elif (new.iloc[i]['action'] == 'Nest') & (new.iloc[i]['reclutch'] != ''):
                    document.add_paragraph('A reclutch individual, carrying tag nos. ' + str(new.iloc[i]['reclutch']) + ', previously nested ' + str(new.iloc[i]['oii']) + ' days ago this season on ' + str(df[df['all.tags'].str.contains(new.iloc[i]['reclutch'][2:6])].tail(3)['nest date'].values[0][0:6]) + ' in ' + str(df[df['all.tags'].str.contains(new.iloc[i]['reclutch'][2:6])].tail(3)['Location'].values[0]) + '. This was her ' + nth[new.iloc[i]['ocf']] + ' clutch of the season. She laid ' + str(new.iloc[i]['eggs']) + ' eggs  which were relocated to the ' + str(new.iloc[i]['hatchery site']) + ' hatchery as nest number ' + str(new.iloc[i]['hatchery nest number']) + '. She had a curved carapace length (CCL) of ' + str(float(new.iloc[i]['ccl'])) + ' and a curved carapace width (CCW) of ' + str(float(new.iloc[i]['ccw'])) + '.', style='List Bullet')

            elif (new.iloc[i]['action'] == 'Nest') & (new.iloc[i]['new.tag'] != '') & (new.iloc[i]['eggs'] != ''):
                    document.add_paragraph('A live turtle nesting was observed on ' + new.iloc[i]['Location'] + '. She was fitted with tag nos. ' + new.iloc[i]['new.tag'] + '. The clutch she deposited contained ' + str(new.iloc[i]['eggs']) + ' eggs which were relocated to the ' + str(new.iloc[i]['hatchery site']) + ' hatchery as nest number ' + str(new.iloc[i]['hatchery nest number']) + '. She had a curved carapace length (CCL) of ' + str(float(new.iloc[i]['ccl'])) + ' and a curved carapace width (CCW) of ' + str(float(new.iloc[i]['ccw'])) + '.', style='List Bullet')

            elif (new.iloc[i]['action'] == 'Nest') & (new.iloc[i]['new.tag'] == '') & (new.iloc[i]['eggs'] != '') & (new.iloc[i]['hatchery site'] != ''):
                    document.add_paragraph('A ' + new.iloc[i]['action'] + ' was found in ' + new.iloc[i]['Location'] + ' containing ' + str(new.iloc[i]['eggs']) + ' eggs. Her clutch was relocated to the ' + str(new.iloc[i]['hatchery site']) + ' hatchery as nest number ' + str(new.iloc[i]['hatchery nest number']) + '.', style='List Bullet')

            elif (new.iloc[i]['action'] == 'Nest') & (new.iloc[i]['new.tag'] == '') & (new.iloc[i]['eggs'] != '') & (new.iloc[i]['hatchery site'] == ''):
                    document.add_paragraph('A ' + new.iloc[i]['action'] + ' was found in ' + new.iloc[i]['Location'] + ' containing ' + str(new.iloc[i]['eggs']) + ' eggs.', style='List Bullet')

            elif (new.iloc[i]['action'] == 'Nest') & (new.iloc[i]['eggs'] == '') & (new.iloc[i]['Location'] != 'Fuwairit'):
                    document.add_paragraph('A nest was found on ' + new.iloc[i]['Location'] + '.', style='List Bullet')

            elif (new.iloc[i]['action'] != 'Nest') & (new.iloc[i]['new.tag'] != ''):
                    document.add_paragraph('A live turtle performing a false crawl was observed on ' + new.iloc[i]['Location'] + '. It was fitted with tag nos. ' + new.iloc[i]['new.tag'] + '. The individual had a curved carapace length (CCL) of ' + str(float(new.iloc[i]['ccl'])) + ' and a curved carapace width (CCW) of ' + str(float(new.iloc[i]['ccw'])) + '.', style='List Bullet')

            elif (new.iloc[i]['action'] != 'Nest') & (new.iloc[i]['new.tag'] == ''):
                    document.add_paragraph('False crawl tracks were observed on ' + new.iloc[i]['Location'] + '.', style='List Bullet')


    #       else:
    #               document.add_paragraph('An ' + new.iloc[i]['action'] + ' was found on ' + new.iloc[i]['Location'] + '.', style='List Bullet')

            i+=1

        im = Image.open(os.path.join('weekly report pics\\' + pics[j]))

        width, height = im.size  # Get dimensions
        new_width = width * 0.70
        new_height = height * 0.50
        left = (width - new_width) / 2
        top = (height - new_height) / 2
        right = (width + new_width) / 2
        bottom = (height + new_height) / 2

        im = im.crop((left, top, right, bottom))
        im.save('test_' + str(j) + '.jpg')
        document.add_picture('test_' + str(j) + '.jpg', width=Inches(4))

        j+=1
        document.add_paragraph('')

    document.add_page_break()

    document.add_heading('Nesting Totals', level=1)

    document.add_paragraph('A total of ' + str(sum(nest_counts['nest date'])) + ' nests were observed during this period of the nesting season while ' + str(sum(fcu_counts['nest date']) + sum(fca_counts['nest date'])) + ' false crawls were recorded. ' + nest_counts['Location'][nest_counts['nest date'].idxmax()] + ' had the highest number of clutches this period with ' + str(nest_counts['nest date'].max()) + '. So far ' + all_nests['Location'][all_nests['nest date'].idxmax()] + ' has recieved the most number of nests for the season with ' + str(all_nests['nest date'].max()) + '. In total ' + str(sum(all_nests['nest date'])) + ' have been recorded this season. The pie chart below shows the relative contribution of each site to the total nesting recorded in the season.')
    document.add_picture('pie.jpg', width=Inches(4.5))
    document.add_paragraph('')

    document.add_paragraph('The figure below shows the frequency of nesting throughout the season.')
    document.add_picture('Nest_freq.jpg', width=Inches(4.5))
    document.add_paragraph('')

    document.add_paragraph('The figure below shows the cumulative nests totalled across all sites.')
    document.add_picture('Cumulative_nests.jpg', width=Inches(4.5))

    menu_table = document.add_table(rows=1, cols=2)
    menu_table.style = 'Table Grid'
    hdr_Cells = menu_table.rows[0].cells
    hdr_Cells[0].text = 'Location'
    hdr_Cells[1].text = 'Total Nests'

    for location, nests in records:
        row_Cells = menu_table.add_row().cells
        row_Cells[0].text = location
        row_Cells[1].text = str(nests)

    outfile = 'QTP_Report_' + str(start_date)[1:11] + '_to_' + str(end_date)[1:11] + '.docx'

    document.save(outfile)



my_btn = Button(root, text='    Generate Report    ', command=generate)
my_btn.pack()


root.mainloop()
